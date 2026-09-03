"""Reporting Engine — generates Markdown and .docx pentest reports."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from .models import Engagement, Finding, Severity
from .scoring import score_from_vector


# ---------------------------------------------------------------------------
# Markdown report
# ---------------------------------------------------------------------------

def _severity_badge(sev: Severity) -> str:
    icons = {
        Severity.CRITICAL: "🔴",
        Severity.HIGH:     "🟠",
        Severity.MEDIUM:   "🟡",
        Severity.LOW:      "🔵",
        Severity.INFO:     "⚪",
    }
    return icons.get(sev, "⚪")


def _deduplicate_findings(engagement: Engagement) -> list[tuple[list[str], Finding]]:
    """Group findings by (tool, title) and return one representative Finding
    per group, alongside every checklist item ID it recurred under.

    Several checklist items commonly run the same tool (nmap alone is now
    mapped to WSTG-INFO-02/04, WSTG-CONF-01/06) — each run's
    extract_findings() independently produces its own Finding object for
    the same real-world fact (e.g. "Server Version Disclosure:
    Apache/httpd"), since findings are tracked per checklist item, not
    deduplicated across the engagement. That's correct for the live
    per-item findings panel (each item genuinely detected it), but it
    means an *un-deduplicated* report shows the identical finding block
    two, three, four times — confirmed against real engagement data
    (one real engagement had "Server Version Disclosure" 4 times).

    Deliberately keys on (tool, title) rather than also including
    `evidence`: the same underlying finding's evidence text can differ
    slightly between runs (`_grep_context()` in findings_extractor.py
    captures a few surrounding lines around the match, which shifts
    depending on where exactly the match lands in that particular
    tool_outputs block) even though it's the same real finding.
    """
    groups: dict[tuple[str, str], list[tuple[str, Finding]]] = {}
    for item in engagement.checklist_items:
        for f in item.findings:
            groups.setdefault((f.tool, f.title), []).append((item.id, f))

    result: list[tuple[list[str], Finding]] = []
    for entries in groups.values():
        item_ids = sorted({item_id for item_id, _ in entries})
        # Prefer a verified finding as the representative if any of the
        # duplicates were manually verified, so a tester's verification
        # work on *any* one instance is what the report reflects.
        representative = next((f for _, f in entries if f.verified), entries[0][1])
        result.append((item_ids, representative))
    return result


# Mirrors frontend/src/lib/methodologies.ts's labels — the executive
# summary should say what methodology actually drove this engagement's
# checklist, not always claim WSTG regardless (a real inaccuracy an OSCP-
# methodology report would otherwise have shipped with).
_METHODOLOGY_LABELS = {
    "wstg": "OWASP WSTG",
    "oscp": "OSCP/PEN-200-style",
    "other": "OWASP WSTG",
}


def _methodology_label(methodology: str) -> str:
    return _METHODOLOGY_LABELS.get(methodology, "OWASP WSTG")


def generate_markdown(engagement: Engagement, out_path: Optional[Path] = None) -> str:
    """Render *engagement* as a Markdown pentest report string."""
    lines: list[str] = []
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    deduped = _deduplicate_findings(engagement)
    sev: dict[str, int] = {s.value: 0 for s in Severity}
    for _item_ids, f in deduped:
        sev[f.severity.value] += 1

    # ── Header ──────────────────────────────────────────────────────────────
    lines += [
        f"# Pentest Report — {engagement.name}",
        "",
        f"| Field | Value |",
        f"|---|---|",
        f"| **Target** | `{engagement.target}` |",
        f"| **Engagement ID** | `{engagement.id}` |",
        f"| **Report Date** | {now} |",
        f"| **Engagement Created** | {engagement.created_at.strftime('%Y-%m-%d %H:%M')} |",
        f"| **Progress** | {engagement.done_items}/{engagement.total_items} checklist items |",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        f"This report presents findings from a deterministic, {_methodology_label(engagement.methodology)} "
        f"checklist-driven assessment of **{engagement.target}**. "
        f"All enumeration was performed using standard open-source tools invoked directly; "
        f"no AI-driven decision-making was used in the enumeration phase.",
        "",
        "### Finding Severity Summary",
        "",
        "| Severity | Count |",
        "|---|---|",
        f"| 🔴 Critical | {sev.get('critical', 0)} |",
        f"| 🟠 High     | {sev.get('high', 0)} |",
        f"| 🟡 Medium   | {sev.get('medium', 0)} |",
        f"| 🔵 Low      | {sev.get('low', 0)} |",
        f"| ⚪ Info     | {sev.get('info', 0)} |",
        f"| **Total**   | **{len(deduped)}** |",
        "",
        "---",
        "",
        "## Checklist Coverage",
        "",
        "| ID | Name | Status | Findings |",
        "|---|---|---|---|",
    ]

    for item in engagement.checklist_items:
        icon = item.status.icon
        count = len(item.findings)
        lines.append(
            f"| `{item.id}` | {item.name} | {icon} {item.status.value} | {count} |"
        )

    lines += [
        "",
        "---",
        "",
        "## Detailed Findings",
        "",
        "*Deduplicated — the same finding detected via more than one checklist "
        "item (e.g. several items running the same tool) is listed once, with "
        "every checklist item it was found under.*",
        "",
    ]

    # Sort: critical → high → medium → low → info
    order = {s.value: i for i, s in enumerate(
        [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]
    )}
    deduped.sort(key=lambda x: order.get(x[1].severity.value, 99))

    if not deduped:
        lines.append("*No findings recorded for this engagement.*")
    else:
        for item_ids, f in deduped:
            badge = _severity_badge(f.severity)
            verified_str = "✅ Verified" if f.verified else "⚠️ Unverified (tool)"
            cvss_str = (
                f"`{f.cvss_vector}` → **{f.cvss_score:.1f}**"
                if f.cvss_vector else f"**{f.cvss_score:.1f}**"
            )
            items_label = "Checklist Items" if len(item_ids) > 1 else "Checklist Item"
            items_value = ", ".join(f"`{i}`" for i in item_ids)
            lines += [
                f"### {badge} [{f.severity.value.upper()}] {f.title}",
                "",
                f"| Field | Value |",
                f"|---|---|",
                f"| **Finding ID** | `{f.id}` |",
                f"| **{items_label}** | {items_value} |",
                f"| **Severity** | {f.severity.value.upper()} |",
                f"| **CVSS Score** | {cvss_str} |",
                f"| **OWASP Category** | {f.owasp_category or '—'} |",
                f"| **CWE** | {f.cwe_id or '—'} |",
                f"| **Tool** | `{f.tool}` |",
                f"| **Status** | {verified_str} |",
                f"| **Discovered** | {f.created_at.strftime('%Y-%m-%d %H:%M')} |",
                "",
                "**Description**",
                "",
                f.description,
                "",
            ]
            if f.evidence:
                lines += [
                    "**Evidence**",
                    "",
                    f"```",
                    f.evidence,
                    "```",
                    "",
                ]
            if f.remediation:
                lines += [
                    "**Remediation**",
                    "",
                    f.remediation,
                    "",
                ]
            lines.append("---")
            lines.append("")

    # Tool output appendix
    lines += ["## Appendix — Raw Tool Output", ""]
    for item in engagement.checklist_items:
        if item.tool_outputs:
            lines.append(f"### {item.id} — {item.name}")
            lines.append("")
            for tool_name, output in item.tool_outputs.items():
                elapsed = item.time_elapsed_seconds
                time_str = f" ({elapsed:.1f}s)" if elapsed else ""
                lines += [
                    f"#### `{tool_name}`{time_str}",
                    "",
                    "```",
                    output.strip(),
                    "```",
                    "",
                ]

    report_text = "\n".join(lines)

    if out_path is not None:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(report_text)

    return report_text


# ---------------------------------------------------------------------------
# Word (.docx) report
# ---------------------------------------------------------------------------

def generate_docx(engagement: Engagement, out_path: Path) -> Path:
    """Generate a .docx report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx is required for .docx export: pip install python-docx"
        )

    doc = Document()

    # Title
    title = doc.add_heading(f"Pentest Report — {engagement.name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta table
    doc.add_heading("Engagement Details", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    meta = [
        ("Target",           engagement.target),
        ("Engagement ID",    engagement.id),
        ("Report Date",      datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Progress",         f"{engagement.done_items}/{engagement.total_items} items"),
        ("Total Findings",   str(len(_deduplicate_findings(engagement)))),
    ]
    for k, v in meta:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    # Findings
    doc.add_heading("Findings", level=1)
    sev_colors = {
        "critical": RGBColor(0xC0, 0x00, 0x00),
        "high":     RGBColor(0xFF, 0x66, 0x00),
        "medium":   RGBColor(0xFF, 0xCC, 0x00),
        "low":      RGBColor(0x00, 0x70, 0xC0),
        "info":     RGBColor(0x70, 0x70, 0x70),
    }

    # Deduplicated — see _deduplicate_findings()'s docstring: several
    # checklist items commonly run the same tool, so the same real finding
    # otherwise appears once per item that detected it.
    deduped = _deduplicate_findings(engagement)

    order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    deduped.sort(key=lambda x: order.get(x[1].severity.value, 99))

    for item_ids, f in deduped:
        heading = doc.add_heading(f"[{f.severity.value.upper()}] {f.title}", level=2)
        run = heading.runs[0]
        run.font.color.rgb = sev_colors.get(f.severity.value, RGBColor(0, 0, 0))

        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Field"
        t.rows[0].cells[1].text = "Value"
        fields = [
            ("Checklist Item" if len(item_ids) == 1 else "Checklist Items", ", ".join(item_ids)),
            ("CVSS Score",     str(f.cvss_score)),
            ("CVSS Vector",    f.cvss_vector or "—"),
            ("OWASP",          f.owasp_category or "—"),
            ("CWE",            f.cwe_id or "—"),
            ("Tool",           f.tool),
            ("Verified",       "Yes" if f.verified else "No (tool output)"),
        ]
        for k, v in fields:
            r = t.add_row().cells
            r[0].text = k
            r[1].text = v

        doc.add_paragraph("Description").bold = True
        doc.add_paragraph(f.description)
        if f.evidence:
            doc.add_paragraph("Evidence").bold = True
            doc.add_paragraph(f.evidence)
        if f.remediation:
            doc.add_paragraph("Remediation").bold = True
            doc.add_paragraph(f.remediation)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
