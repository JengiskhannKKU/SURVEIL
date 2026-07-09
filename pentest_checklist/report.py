"""Reporting Engine — generates Markdown and .docx pentest reports."""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Optional

from .models import Engagement, Finding, Severity
from .scoring import score_from_vector


# ---------------------------------------------------------------------------
# Markdown report
# ---------------------------------------------------------------------------

def _severity_badge(sev: Severity) -> str:
    icons = {
        Severity.CRITICAL: "🔴",
        Severity.HIGH:     "🟠",
        Severity.MEDIUM:   "🟡",
        Severity.LOW:      "🔵",
        Severity.INFO:     "⚪",
    }
    return icons.get(sev, "⚪")


def generate_markdown(engagement: Engagement, out_path: Optional[Path] = None) -> str:
    """Render *engagement* as a Markdown pentest report string."""
    lines: list[str] = []
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    sev = engagement.findings_by_severity

    # ── Header ──────────────────────────────────────────────────────────────
    lines += [
        f"# Pentest Report — {engagement.name}",
        "",
        f"| Field | Value |",
        f"|---|---|",
        f"| **Target** | `{engagement.target}` |",
        f"| **Engagement ID** | `{engagement.id}` |",
        f"| **Report Date** | {now} |",
        f"| **Engagement Created** | {engagement.created_at.strftime('%Y-%m-%d %H:%M')} |",
        f"| **Progress** | {engagement.done_items}/{engagement.total_items} checklist items |",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        f"This report presents findings from a deterministic, OWASP WSTG checklist-driven "
        f"web application security assessment of **{engagement.target}**. "
        f"All enumeration was performed using standard open-source tools invoked directly; "
        f"no AI-driven decision-making was used in the enumeration phase.",
        "",
        "### Finding Severity Summary",
        "",
        "| Severity | Count |",
        "|---|---|",
        f"| 🔴 Critical | {sev.get('critical', 0)} |",
        f"| 🟠 High     | {sev.get('high', 0)} |",
        f"| 🟡 Medium   | {sev.get('medium', 0)} |",
        f"| 🔵 Low      | {sev.get('low', 0)} |",
        f"| ⚪ Info     | {sev.get('info', 0)} |",
        f"| **Total**   | **{engagement.total_findings}** |",
        "",
        "---",
        "",
        "## Checklist Coverage",
        "",
        "| ID | Name | Status | Findings |",
        "|---|---|---|---|",
    ]

    for item in engagement.checklist_items:
        icon = item.status.icon
        count = len(item.findings)
        lines.append(
            f"| `{item.id}` | {item.name} | {icon} {item.status.value} | {count} |"
        )

    lines += ["", "---", "", "## Detailed Findings", ""]

    # Group all findings by severity
    all_findings: list[tuple[str, Finding]] = []
    for item in engagement.checklist_items:
        for f in item.findings:
            all_findings.append((item.id, f))

    # Sort: critical → high → medium → low → info
    order = {s.value: i for i, s in enumerate(
        [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]
    )}
    all_findings.sort(key=lambda x: order.get(x[1].severity.value, 99))

    if not all_findings:
        lines.append("*No findings recorded for this engagement.*")
    else:
        for item_id, f in all_findings:
            badge = _severity_badge(f.severity)
            verified_str = "✅ Verified" if f.verified else "⚠️ Unverified (tool)"
            cvss_str = (
                f"`{f.cvss_vector}` → **{f.cvss_score:.1f}**"
                if f.cvss_vector else f"**{f.cvss_score:.1f}**"
            )
            lines += [
                f"### {badge} [{f.severity.value.upper()}] {f.title}",
                "",
                f"| Field | Value |",
                f"|---|---|",
                f"| **Finding ID** | `{f.id}` |",
                f"| **Checklist Item** | `{item_id}` |",
                f"| **Severity** | {f.severity.value.upper()} |",
                f"| **CVSS Score** | {cvss_str} |",
                f"| **OWASP Category** | {f.owasp_category or '—'} |",
                f"| **CWE** | {f.cwe_id or '—'} |",
                f"| **Tool** | `{f.tool}` |",
                f"| **Status** | {verified_str} |",
                f"| **Discovered** | {f.created_at.strftime('%Y-%m-%d %H:%M')} |",
                "",
                "**Description**",
                "",
                f.description,
                "",
            ]
            if f.evidence:
                lines += [
                    "**Evidence**",
                    "",
                    f"```",
                    f.evidence,
                    "```",
                    "",
                ]
            if f.remediation:
                lines += [
                    "**Remediation**",
                    "",
                    f.remediation,
                    "",
                ]
            lines.append("---")
            lines.append("")

    # Tool output appendix
    lines += ["## Appendix — Raw Tool Output", ""]
    for item in engagement.checklist_items:
        if item.tool_outputs:
            lines.append(f"### {item.id} — {item.name}")
            lines.append("")
            for tool_name, output in item.tool_outputs.items():
                elapsed = item.time_elapsed_seconds
                time_str = f" ({elapsed:.1f}s)" if elapsed else ""
                lines += [
                    f"#### `{tool_name}`{time_str}",
                    "",
                    "```",
                    output.strip(),
                    "```",
                    "",
                ]

    report_text = "\n".join(lines)

    if out_path is not None:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(report_text)

    return report_text


# ---------------------------------------------------------------------------
# Word (.docx) report
# ---------------------------------------------------------------------------

def generate_docx(engagement: Engagement, out_path: Path) -> Path:
    """Generate a .docx report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError(
            "python-docx is required for .docx export: pip install python-docx"
        )

    doc = Document()

    # Title
    title = doc.add_heading(f"Pentest Report — {engagement.name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta table
    doc.add_heading("Engagement Details", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    meta = [
        ("Target",           engagement.target),
        ("Engagement ID",    engagement.id),
        ("Report Date",      datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Progress",         f"{engagement.done_items}/{engagement.total_items} items"),
        ("Total Findings",   str(engagement.total_findings)),
    ]
    for k, v in meta:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = v

    # Findings
    doc.add_heading("Findings", level=1)
    sev_colors = {
        "critical": RGBColor(0xC0, 0x00, 0x00),
        "high":     RGBColor(0xFF, 0x66, 0x00),
        "medium":   RGBColor(0xFF, 0xCC, 0x00),
        "low":      RGBColor(0x00, 0x70, 0xC0),
        "info":     RGBColor(0x70, 0x70, 0x70),
    }

    all_findings: list[tuple[str, Finding]] = []
    for item in engagement.checklist_items:
        for f in item.findings:
            all_findings.append((item.id, f))

    order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    all_findings.sort(key=lambda x: order.get(x[1].severity.value, 99))

    for item_id, f in all_findings:
        heading = doc.add_heading(f"[{f.severity.value.upper()}] {f.title}", level=2)
        run = heading.runs[0]
        run.font.color.rgb = sev_colors.get(f.severity.value, RGBColor(0, 0, 0))

        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Field"
        t.rows[0].cells[1].text = "Value"
        fields = [
            ("Checklist Item", item_id),
            ("CVSS Score",     str(f.cvss_score)),
            ("CVSS Vector",    f.cvss_vector or "—"),
            ("OWASP",          f.owasp_category or "—"),
            ("CWE",            f.cwe_id or "—"),
            ("Tool",           f.tool),
            ("Verified",       "Yes" if f.verified else "No (tool output)"),
        ]
        for k, v in fields:
            r = t.add_row().cells
            r[0].text = k
            r[1].text = v

        doc.add_paragraph("Description").bold = True
        doc.add_paragraph(f.description)
        if f.evidence:
            doc.add_paragraph("Evidence").bold = True
            doc.add_paragraph(f.evidence)
        if f.remediation:
            doc.add_paragraph("Remediation").bold = True
            doc.add_paragraph(f.remediation)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
